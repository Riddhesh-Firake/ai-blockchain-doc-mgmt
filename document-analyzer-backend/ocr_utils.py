"""
OCR and text extraction utilities for various document formats.
Supports PDF (including scanned), DOCX, DOC, and TXT files.
"""

import os
import logging
from typing import Optional
from pathlib import Path

# Text extraction libraries
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import fitz  # PyMuPDF for PDF text extraction
from docx import Document
import docx2txt

# Configure logging
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using multiple methods:
    1. Try direct text extraction (for text-based PDFs)
    2. Fall back to OCR (for scanned PDFs)
    """
    text = ""
    
    try:
        # Method 1: Try direct text extraction using PyMuPDF
        logger.info("Attempting direct text extraction from PDF...")
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text()
            text += page_text
        
        doc.close()
        
        # If we got good text content, return it
        if len(text.strip()) > 100:  # Arbitrary threshold for "good" text
            logger.info(f"Direct extraction successful: {len(text)} characters")
            return text.strip()
        
        logger.info("Direct extraction yielded minimal text, trying OCR...")
        
    except Exception as e:
        logger.warning(f"Direct PDF text extraction failed: {e}, trying OCR...")
    
    try:
        # Method 2: OCR approach using pdf2image + pytesseract
        logger.info("Converting PDF to images for OCR...")
        
        # Convert PDF to images (one per page)
        pages = convert_from_path(
            file_path,
            dpi=300,  # Higher DPI for better OCR accuracy
            first_page=1,
            last_page=None,
            poppler_path=None  # Use system poppler if available
        )
        
        ocr_text = ""
        for i, page in enumerate(pages):
            logger.info(f"Performing OCR on page {i+1}/{len(pages)}...")
            
            # Use pytesseract to extract text from image
            page_text = pytesseract.image_to_string(
                page,
                config='--oem 3 --psm 6'  # OCR Engine Mode 3, Page Segmentation Mode 6
            )
            ocr_text += page_text + "\n"
        
        if ocr_text.strip():
            logger.info(f"OCR extraction successful: {len(ocr_text)} characters")
            return ocr_text.strip()
        
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        raise Exception(f"Failed to extract text from PDF: {e}")
    
    raise Exception("Could not extract any text from PDF file")

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX files using python-docx.
    """
    try:
        logger.info("Extracting text from DOCX file...")
        
        # Method 1: Use docx2txt (simple and effective)
        text = docx2txt.process(file_path)
        
        if text and len(text.strip()) > 0:
            logger.info(f"DOCX extraction successful: {len(text)} characters")
            return text.strip()
        
        # Method 2: Use python-docx for more control
        doc = Document(file_path)
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        text = "\n".join(paragraphs)
        
        if text.strip():
            logger.info(f"DOCX extraction (method 2) successful: {len(text)} characters")
            return text.strip()
        
        raise Exception("No text content found in DOCX file")
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise Exception(f"Failed to extract text from DOCX: {e}")

def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from legacy DOC files.
    This is more challenging and may require additional tools.
    """
    try:
        logger.info("Extracting text from DOC file...")
        
        # For DOC files, we could use python-docx2txt or antiword
        # For now, let's try a simple approach and suggest conversion
        raise Exception("DOC file support requires additional tools. Please convert to DOCX format.")
        
    except Exception as e:
        logger.error(f"DOC extraction failed: {e}")
        raise Exception(f"Failed to extract text from DOC file: {e}")

def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from plain text files with encoding detection.
    """
    try:
        logger.info("Reading text file...")
        
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    if text.strip():
                        logger.info(f"Text file read successfully with {encoding}: {len(text)} characters")
                        return text.strip()
            except UnicodeDecodeError:
                continue
        
        raise Exception("Could not decode text file with any supported encoding")
        
    except Exception as e:
        logger.error(f"Text file reading failed: {e}")
        raise Exception(f"Failed to read text file: {e}")

def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """
    Main function to extract text from various file formats.
    
    Args:
        file_path: Path to the file
        file_extension: File extension (e.g., '.pdf', '.docx')
    
    Returns:
        Extracted text content
    """
    file_extension = file_extension.lower()
    
    logger.info(f"Extracting text from file: {Path(file_path).name} (extension: {file_extension})")
    
    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        elif file_extension == '.doc':
            return extract_text_from_doc(file_path)
        elif file_extension == '.txt':
            return extract_text_from_txt(file_path)
        else:
            raise Exception(f"Unsupported file extension: {file_extension}")
    
    except Exception as e:
        logger.error(f"Text extraction failed for {file_extension} file: {e}")
        raise

def preprocess_text(text: str) -> str:
    """
    Clean and preprocess extracted text for better NLP analysis.
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if line:  # Only keep non-empty lines
            cleaned_lines.append(line)
    
    # Join with single newlines
    cleaned_text = '\n'.join(cleaned_lines)
    
    # Remove excessive spaces
    import re
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
    
    return cleaned_text.strip()

# Function to validate if pytesseract is properly configured
def check_ocr_dependencies():
    """
    Check if OCR dependencies are properly installed and configured.
    """
    try:
        # Test pytesseract
        pytesseract.get_tesseract_version()
        logger.info("Tesseract OCR is properly configured")
        return True
    except Exception as e:
        logger.error(f"OCR dependencies not properly configured: {e}")
        logger.error("Please install Tesseract OCR and ensure it's in your PATH")
        return False

# Initialize OCR check on module import
if __name__ != "__main__":
    check_ocr_dependencies()